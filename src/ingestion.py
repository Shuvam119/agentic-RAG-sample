"""Document loading, metadata extraction, and chunking.

This project builds its own index from the raw documents in DOCS_SOURCE
(default: the Traditional RAG project's data/raw folder). The pipeline is
self-contained and does not touch the Traditional RAG project's index.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from src.config import CHUNK_OVERLAP, CHUNK_SIZE, DOCS_SOURCE, SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------

def load_docx(file_path: str) -> Dict:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    doc = Document(path)
    paragraphs = [
        para.text for para in doc.paragraphs if para.text.strip()
    ]

    for table in doc.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if cells:
                paragraphs.append(" | ".join(cells))

    full_text = "\n".join(paragraphs)

    return {
        "title": path.name,
        "text": full_text,
        "metadata": {
            "source": str(path),
            "type": "docx",
            "filename": path.name,
            "empty_text": not full_text.strip(),
        },
    }


def load_pdf(file_path: str) -> Dict:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    reader = PdfReader(str(path))
    pages = []

    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {page_num} ---\n{text}")

    full_text = "\n".join(pages)

    return {
        "title": path.name,
        "text": full_text,
        "metadata": {
            "source": str(path),
            "type": "pdf",
            "filename": path.name,
            "pages": len(reader.pages),
            "empty_text": not full_text.strip(),
        },
    }


def load_documents_from_directory(directory: str | Path | None = None) -> List[Dict]:
    directory = Path(directory or DOCS_SOURCE)

    if not directory.is_dir():
        raise ValueError(f"Not a directory: {directory}")

    documents: List[Dict] = []

    for file_path in sorted(directory.iterdir()):
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        loader = load_docx if file_path.suffix.lower() == ".docx" else load_pdf

        try:
            documents.append(loader(str(file_path)))
        except Exception as exc:
            print(f"Error loading {file_path.name}: {exc}")

    return documents


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------

DOCUMENT_TYPE_ALIASES = {
    "standard operating procedure": "SOP",
    "sop": "SOP",
    "user guide": "User Guide",
    "users guide": "User Guide",
    "user's guide": "User Guide",
    "administrator guide": "Administrator Guide",
    "administrators guide": "Administrator Guide",
    "admin guide": "Administrator Guide",
    "job aid": "Job Aid",
    "quick reference": "Job Aid",
    "quick reference guide": "Job Aid",
    "faq": "FAQ",
    "release notes": "Release Notes",
    "release note": "Release Notes",
    "known issues": "Known Issues",
    "known issue": "Known Issues",
    "api documentation": "API Documentation",
    "api guide": "API Documentation",
    "api reference": "API Documentation",
    "api reference guide": "API Documentation",
    "training manual": "Training Manual",
    "training guide": "Training Manual",
    "troubleshooting guide": "Troubleshooting Guide",
    "troubleshooting": "Troubleshooting Guide",
}

_FILENAME_TYPE_SIGNALS = (
    ("api", "API Documentation"),
    ("integration guide", "API Documentation"),
    ("job aid", "Job Aid"),
    ("quick reference", "Job Aid"),
    ("sop", "SOP"),
    ("faq", "FAQ"),
    ("release notes", "Release Notes"),
    ("known issues", "Known Issues"),
    ("user guide", "User Guide"),
    ("users guide", "User Guide"),
    ("admin guide", "Administrator Guide"),
    ("administrator", "Administrator Guide"),
    ("training", "Training Manual"),
    ("troubleshoot", "Troubleshooting Guide"),
)

_HEADER_BODY_MARKERS = (
    r"summary",
    r"table\s+of\s+contents",
    r"\d+\.\s+(?:introduction|overview|purpose\s+and\s+scope)",
)

_STOP_WORDS = {
    "this", "that", "with", "from", "your", "will", "have", "into",
    "using", "when", "where", "document", "guide", "version",
    "the", "and", "for", "are", "you",
}

LIFECYCLE_STATUSES = (
    "Fresh", "Need Update", "Needs Deprecation", "Aging", "Stale",
    "Archived", "Needs Review",
)


def _title_from_filename(filename: str) -> str:
    stem = Path(filename).stem.replace("_", " ").replace("-", " ")
    return re.sub(r"\b(v(?:ersion)?\s*\d+(?:\.\d+)*)\b", "", stem, flags=re.I).strip().title()


def _first_match(pattern: str, text: str, default: str) -> str:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(1).strip() if match else default


def _header_section(text: str, max_chars: int = 3000) -> str:
    marker = re.search(
        r"(?im)^\s*(?:%s)\b" % "|".join(_HEADER_BODY_MARKERS), text)
    if marker:
        text = text[:marker.start()]
    return text[:max_chars]


def _keywords(text: str, limit: int = 10) -> List[str]:
    words = re.findall(r"[A-Za-z][A-Za-z0-9_-]{3,}", text.lower())
    unique: List[str] = []
    for word in words:
        if word not in _STOP_WORDS and word not in unique:
            unique.append(word)
        if len(unique) == limit:
            break
    return unique


def _parse_iso_date(value: str) -> date | None:
    try:
        return date.fromisoformat(value)
    except Exception:
        return None


def _normalize_document_type(value: str) -> str:
    cleaned = re.sub(r"\s*\|\s*.*$", "", value).strip()
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    lowered = cleaned.lower()

    # PDF column layouts often append extra status text after the type
    # (e.g. "User Guide Approval Status Approved"). If the value starts
    # with a known type, use it.
    best = None
    for alias, kind in DOCUMENT_TYPE_ALIASES.items():
        if lowered.startswith(alias) and (
            best is None or len(alias) > len(best[0])
        ):
            best = (alias, kind)
    if best:
        return best[1]

    return DOCUMENT_TYPE_ALIASES.get(lowered, cleaned)


def _declared_document_type(text: str) -> str:
    match = re.search(
        r"(?im)^\s*document\s+type\b\s*(?:[:|\-])?\s*([^\n|]{1,80})", text)
    return _normalize_document_type(match.group(1)) if match else ""


def _infer_document_type(haystack: str) -> str:
    filename = haystack.split("\n", 1)[0].lower()
    for signal, kind in _FILENAME_TYPE_SIGNALS:
        if signal in filename:
            return kind
    lowered = haystack.lower()
    return next(
        (kind for kind in DOCUMENT_TYPE_ALIASES.values() if kind.lower() in lowered),
        "User Guide",
    )


def _document_type_from_document(text: str, filename: str) -> str:
    declared = _declared_document_type(text)
    if declared:
        return declared
    return _infer_document_type(f"{filename}\n{text[:5000]}")


def extract_metadata(document: Dict[str, Any]) -> Dict[str, Any]:
    source = document["metadata"]
    text = document.get("text", "")
    filename = source["filename"]
    haystack = f"{filename}\n{text[:5000]}"
    lowered = haystack.lower()

    document_type = _document_type_from_document(text, filename)

    audience = "End User"
    if "administrator" in lowered or "admin guide" in lowered:
        audience = "Administrator"
    elif "support engineer" in lowered or "troubleshooting" in lowered:
        audience = "Support Engineer"
    elif "technical writer" in lowered:
        audience = "Technical Writer"
    elif "product manager" in lowered:
        audience = "Product Manager"

    version = _first_match(
        r"\b(?:version|ver|v)\s*(\d+(?:\.\d+)*)\b", haystack, "Unspecified")
    product = _first_match(
        r"(?:product|application|platform)\s*[:\-]\s*([^\n]{2,60})",
        _header_section(text),
        "General",
    )
    department = _first_match(
        r"(?:department|owner|team)\s*[:\-]\s*([^\n]{2,60})", text, "Documentation")
    author = _first_match(
        r"(?:author|written by|created by)\s*[:\-]\s*([^\n]{2,80})", text, "Unknown")
    publication_date = _first_match(
        r"(?:published|publication|released)(?:\s+on)?\s*[:\-]?\s*([0-9]{4}-[0-9]{2}-[0-9]{2})",
        text,
        "",
    )
    summary_source = " ".join(text.split())
    summary = summary_source[:360].rsplit(" ", 1)[0] if len(
        summary_source) > 360 else summary_source

    return {
        "title": _title_from_filename(filename),
        "product": product,
        "version": version,
        "document_type": document_type,
        "audience": audience,
        "department": department,
        "author": author,
        "last_updated": str(date.today()),
        "publication_date": publication_date or str(date.today()),
        "keywords": _keywords(haystack),
        "summary": summary or "No extractable summary.",
        "lifecycle_status": determine_lifecycle_status(
            {"version": version, "title": _title_from_filename(filename),
             "document_type": document_type, "last_updated": str(date.today())}
        ),
    }


def determine_lifecycle_status(metadata: Dict[str, Any]) -> str:
    explicit = metadata.get("lifecycle_status")
    if explicit in LIFECYCLE_STATUSES:
        return explicit

    version = str(metadata.get("version", "Unspecified")).strip().lower()
    if version == "unspecified":
        return "Fresh"

    title = str(metadata.get("title", "")).lower()
    if "archive" in title or "legacy" in title:
        return "Archived"

    last_updated = _parse_iso_date(metadata.get("last_updated", ""))
    publication_date = _parse_iso_date(metadata.get("publication_date", ""))
    reference_date = last_updated or publication_date or date.today()
    days_old = (date.today() - reference_date).days

    if days_old <= 60:
        return "Fresh"
    if days_old <= 180:
        return "Aging"
    if days_old <= 365:
        return "Needs Review"
    return "Stale"


def parse_version_tuple(version: str) -> tuple[int, ...] | None:
    match = re.search(r"(\d+(?:\.\d+)*)", str(version))
    if not match:
        return None
    return tuple(int(part) for part in match.group(1).split('.'))


def _normalized_version(version: tuple[int, ...], length: int = 3) -> tuple[int, ...]:
    return version + (0,) * max(0, length - len(version))


def apply_version_lifecycle(documents: List[Dict]) -> None:
    """Latest version of a title is Fresh; older versions become Need Update
    or Needs Deprecation."""
    groups: Dict[str, List[Dict]] = {}
    for document in documents:
        title = document["metadata"].get("title", "")
        groups.setdefault(title, []).append(document)

    for group in groups.values():
        versioned = [
            (document, parse_version_tuple(document["metadata"].get("version", "")))
            for document in group
        ]
        available = [version for _, version in versioned if version is not None]
        if not available:
            continue

        latest = max(available, key=_normalized_version)
        latest_normalized = _normalized_version(latest)

        for document, version in versioned:
            if version is None:
                continue
            normalized = _normalized_version(version)
            if normalized == latest_normalized:
                document["metadata"]["lifecycle_status"] = "Fresh"
            elif latest_normalized[0] - normalized[0] >= 2:
                document["metadata"]["lifecycle_status"] = "Needs Deprecation"
            else:
                document["metadata"]["lifecycle_status"] = "Need Update"


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

class DocumentChunker:
    def __init__(self, chunk_size: int = CHUNK_SIZE,
                 chunk_overlap: int = CHUNK_OVERLAP) -> None:
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def chunk_documents(self, documents: List[Dict]) -> List[Dict]:
        chunks: List[Dict] = []

        for document in documents:
            text = document["text"]
            metadata = document["metadata"]

            chunk_texts = self.splitter.split_text(text)
            total_chunks = len(chunk_texts)
            filename = metadata["filename"]

            for chunk_idx, chunk_text in enumerate(chunk_texts, 1):
                chunk_id = f"{filename}_{chunk_idx}"
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        **metadata,
                        "filename": filename,
                        "chunk_id": chunk_id,
                        "chunk_index": chunk_idx,
                        "chunk_size": len(chunk_text),
                        "total_chunks": total_chunks,
                    },
                })

        return chunks
